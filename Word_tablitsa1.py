# -*- coding: utf-8 -*-
import docx
from docx import Document
from docx import table
import os
import datascience

doc = docx.Document()

doc.add_heading('Вхідні дані(Таблиця 1)', 0)

enterprise = (
      ['Універсам 22', '1', '44 048,00', '500,00', '200,00'],
      ['Дніпрянка   ', '1', '22 456,00', '365,00', '123,00'],
      ['Універсам 22', '2', '5 564,00', '2 751,00', '135,00'],
      ['Дніпрянка   ', '2', '10 087,00', '15 972,00', '135,00'],
      ['Універсам 22', '3', '0,00', '0,00', '0,00'],
      ['Дніпрянка   ', '3', '206,00', '34,00', '50,00'],
      ['Універсам 22', '4', '116,00', '64,00', '23,00'],
      ['Дніпрянка   ', '4', '34,00', '23,00', '25,00'],
)

menuTable = doc.add_table(rows = 1, cols = 5, style = 'TableGrid')
hdr_cells = table.Row[0].cells
hdr_cells[0].text = 'Підприємство'
hdr_cells[1].text = 'Код'
hdr_cells[2].text = 'Залишок на 1.01.2018'
hdr_cells[3].text = 'Надійшло у 2018'
hdr_cells[4].text = 'Вибуток у 2018'

for name, code, remainder, received, output in enterprise:
    row_Cells = table.add_Row().cells
    row_Cells[0].text = name
    row_Cells[1].text = code
    row_Cells[2].text = remainder
    row_Cells[3].text = received
    row_Cells[4].text = output

doc.save('Word_tablitsa1.docx')
