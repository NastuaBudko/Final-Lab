# -*- coding: utf-8 -*-
import docx
from docx import Document
from docx import table
import os
import datascience

doc = docx.Document()

doc.add_heading('Вхідні дані(Таблиця 2)', 0)

enterprise = (
            ['1', 'Будівлі',
            '2','Машини та обладнання',
            '3', 'Транспортні засоби',
            '4', 'Інструмент та інвентар',]
)

menuTable = doc.add_table(rows = 2, cols = 2, style = 'TableGrid')
hdr_cells = table.Row[0].cells
hdr_cells[0].text = 'Код'
hdr_cells[1].text = 'Від основних засобів'

for name, code, remainder, received, output in enterprise:
    row_Cells = table.add_Row().cells
    row_Cells[0].text = code
    row_Cells[1].text = type


doc.save('Word_tablitsa2.docx')
